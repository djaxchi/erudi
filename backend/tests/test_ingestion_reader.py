"""P4 — DocumentReader façade + Tier-0 extractors + text cleaning.

The reader is the ONLY surface services see: it routes by document type
(extension), then by difficulty for PDFs (text layer → deterministic
fast-path; scanned/no-text → ``pending_vision``, indexed by the OCR/VLM
tiers of a later release). Images are accepted but ``pending_vision``.

Cleaning is non-destructive: NFC normalization, control-char removal,
whitespace collapsing, PDF hyphenation rejoin — accents and non-ASCII are
PRESERVED (the FAISS-era cleaner stripped them, killing multilingual
retrieval).
"""

import pytest

from src.core.exceptions import InvalidInputException
from src.ingestion.cleaning import clean_extracted_text
from src.ingestion.reader import DocumentReader

pytestmark = pytest.mark.unit


# ---------------------------------------------------------------------------
# Fixture builders — real files, real parsers, no mocks.
# ---------------------------------------------------------------------------

def _escape_pdf_text(text: str) -> str:
    return text.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")


def _minimal_pdf(pages_text: list[str]) -> bytes:
    """Hand-built single-font PDF with one text line per page.

    Uses WinAnsiEncoding so latin-1 accented characters (é, à, ï…) survive
    the trip through pypdf's extractor.
    """
    objects: list[bytes] = []

    n_pages = len(pages_text)
    first_page_obj = 4
    page_objs = [first_page_obj + 2 * i for i in range(n_pages)]
    kids = " ".join(f"{n} 0 R" for n in page_objs)

    objects.append(b"<< /Type /Catalog /Pages 2 0 R >>")  # 1
    objects.append(
        f"<< /Type /Pages /Kids [{kids}] /Count {n_pages} >>".encode()
    )  # 2
    objects.append(
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica "
        b"/Encoding /WinAnsiEncoding >>"
    )  # 3

    for i, text in enumerate(pages_text):
        content = (
            f"BT /F1 12 Tf 72 720 Td ({_escape_pdf_text(text)}) Tj ET"
        ).encode("latin-1")
        page = (
            f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            f"/Resources << /Font << /F1 3 0 R >> >> "
            f"/Contents {page_objs[i] + 1} 0 R >>"
        ).encode()
        objects.append(page)
        objects.append(
            f"<< /Length {len(content)} >>\nstream\n".encode()
            + content
            + b"\nendstream"
        )

    out = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for n, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{n} 0 obj\n".encode() + body + b"\nendobj\n"
    xref_at = len(out)
    out += f"xref\n0 {len(objects) + 1}\n".encode()
    out += b"0000000000 65535 f \n"
    for off in offsets[1:]:
        out += f"{off:010d} 00000 n \n".encode()
    out += (
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_at}\n%%EOF\n"
    ).encode()
    return bytes(out)


@pytest.fixture
def pdf_with_text(tmp_path):
    path = tmp_path / "report.pdf"
    path.write_bytes(
        _minimal_pdf(
            [
                "Le café des développeurs ouvre ses portes à toute l'équipe.",
                "Deuxième page très utile avec encore plus de contenu textuel.",
            ]
        )
    )
    return path


@pytest.fixture
def pdf_scanned(tmp_path):
    """A PDF whose pages carry no text layer (a scan, as pypdf sees it)."""
    from pypdf import PdfWriter

    path = tmp_path / "scan.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    writer.add_blank_page(width=612, height=792)
    with path.open("wb") as fh:
        writer.write(fh)
    return path


@pytest.fixture
def docx_file(tmp_path):
    import docx

    path = tmp_path / "memo.docx"
    document = docx.Document()
    document.add_heading("Présentation générale", level=1)
    document.add_paragraph("Résumé à l'attention des employés.")
    document.add_heading("Détails", level=2)
    document.add_paragraph("Les données ci-dessous récapitulent l'année.")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Trimestre"
    table.rows[0].cells[1].text = "Chiffre"
    table.rows[1].cells[0].text = "T1"
    table.rows[1].cells[1].text = "1 200 €"
    document.save(str(path))
    return path


@pytest.fixture
def xlsx_file(tmp_path):
    import openpyxl

    path = tmp_path / "budget.xlsx"
    wb = openpyxl.Workbook()
    ws1 = wb.active
    ws1.title = "Dépenses"
    ws1.append(["Poste", "Montant"])
    ws1.append(["Café", 42.5])
    ws2 = wb.create_sheet("Synthèse")
    ws2.append(["Total", 42.5])
    wb.save(str(path))
    return path


@pytest.fixture
def csv_file(tmp_path):
    path = tmp_path / "données.csv"
    path.write_text("nom,rôle\nHélène,ingénieure\nJosé,développeur\n", encoding="utf-8")
    return path


# ---------------------------------------------------------------------------
# Cleaning — non-destructive (D4)
# ---------------------------------------------------------------------------

class TestCleaning:
    def test_accents_and_non_ascii_preserved(self):
        text = "Le café naïve coûte 5 € — vraiment ?"
        assert clean_extracted_text(text) == text

    def test_nfc_normalization(self):
        # e + combining acute accent (U+0301) → precomposed é (U+00E9)
        assert clean_extracted_text("cafe\u0301") == "caf\u00e9"

    def test_pdf_hyphenation_rejoined(self):
        assert clean_extracted_text("informa-\ntion utile") == "information utile"

    def test_control_chars_removed_newlines_kept_tabs_normalized(self):
        # NUL and form-feed dropped; newlines kept; tabs are horizontal
        # whitespace → collapsed to a single space.
        assert clean_extracted_text("a\x00b\x0cc\nd\te") == "abc\nd e"

    def test_whitespace_collapsed_paragraphs_kept(self):
        assert clean_extracted_text("a   b\n\n\n\nc  d") == "a b\n\nc d"

    def test_empty_input(self):
        assert clean_extracted_text("") == ""


# ---------------------------------------------------------------------------
# Reader routing + per-format extraction
# ---------------------------------------------------------------------------

class TestPdf:
    def test_text_layer_pdf_is_active_with_pages(self, pdf_with_text):
        doc = DocumentReader().read(pdf_with_text)
        assert doc.status == "active"
        assert "café" in doc.markdown  # accents survive extraction [M4]
        assert "Deuxième page" in doc.markdown
        assert doc.pages is not None and len(doc.pages) == 2
        assert [p.page_number for p in doc.pages] == [1, 2]
        assert doc.metadata["page_count"] == 2

    def test_scanned_pdf_goes_pending_vision(self, pdf_scanned):
        doc = DocumentReader().read(pdf_scanned)
        assert doc.status == "pending_vision"
        assert doc.markdown == ""


class TestDocx:
    def test_headings_become_markdown_and_text_extracted(self, docx_file):
        doc = DocumentReader().read(docx_file)
        assert doc.status == "active"
        assert "# Présentation générale" in doc.markdown
        assert "## Détails" in doc.markdown
        assert "Résumé à l'attention des employés." in doc.markdown

    def test_tables_become_markdown_tables(self, docx_file):
        doc = DocumentReader().read(docx_file)
        assert "| Trimestre | Chiffre |" in doc.markdown
        assert "| T1 | 1 200 € |" in doc.markdown


class TestXlsx:
    def test_each_sheet_becomes_a_section_with_table(self, xlsx_file):
        doc = DocumentReader().read(xlsx_file)
        assert doc.status == "active"
        assert "## Dépenses" in doc.markdown
        assert "## Synthèse" in doc.markdown
        assert "| Poste | Montant |" in doc.markdown
        assert "| Café | 42.5 |" in doc.markdown


class TestCsv:
    def test_csv_becomes_markdown_table(self, csv_file):
        doc = DocumentReader().read(csv_file)
        assert doc.status == "active"
        assert "| nom | rôle |" in doc.markdown
        assert "| Hélène | ingénieure |" in doc.markdown


class TestText:
    def test_txt_cleaned_passthrough(self, tmp_path):
        path = tmp_path / "notes.txt"
        path.write_text("Première   ligne avec   espaces\n\n\n\nDeuxième.", encoding="utf-8")
        doc = DocumentReader().read(path)
        assert doc.status == "active"
        assert doc.markdown == "Première ligne avec espaces\n\nDeuxième."

    def test_md_keeps_structure(self, tmp_path):
        path = tmp_path / "guide.md"
        path.write_text("# Titre\n\nDu texte **gras** é.", encoding="utf-8")
        doc = DocumentReader().read(path)
        assert "# Titre" in doc.markdown
        assert "é" in doc.markdown

    def test_latin1_fallback(self, tmp_path):
        path = tmp_path / "legacy.txt"
        path.write_bytes("café légataire".encode("latin-1"))
        doc = DocumentReader().read(path)
        assert "café légataire" in doc.markdown


class TestRouting:
    def test_image_is_accepted_pending_vision(self, tmp_path):
        path = tmp_path / "photo.jpg"
        path.write_bytes(b"\xff\xd8\xff\xe0 not a real jpeg")
        doc = DocumentReader().read(path)
        assert doc.status == "pending_vision"
        assert doc.markdown == ""

    @pytest.mark.parametrize("ext", [".png", ".jpeg", ".webp"])
    def test_all_image_extensions_pending_vision(self, tmp_path, ext):
        path = tmp_path / f"img{ext}"
        path.write_bytes(b"fake")
        assert DocumentReader().read(path).status == "pending_vision"

    def test_unsupported_extension_raises(self, tmp_path):
        path = tmp_path / "data.xyz"
        path.write_text("whatever")
        with pytest.raises(InvalidInputException):
            DocumentReader().read(path)

    def test_extension_routing_is_case_insensitive(self, tmp_path):
        path = tmp_path / "NOTES.TXT"
        path.write_text("contenu", encoding="utf-8")
        assert DocumentReader().read(path).status == "active"
