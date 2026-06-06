"""Generate the "Nimbus Analytics" reference KB corpus for RAG quality evals.

Six realistic, fact-dense documents about a fictional French SaaS company.
Every fact (price, SLA, date, headcount…) is PLANTED and cross-checkable, so
a judge can verify any answer against the source. Formats cover the whole
Tier-0 extraction surface: md, txt-equivalent FAQ (md), docx (headings +
tables), xlsx (multi-sheet), and a multi-page text-layer pdf.

Usage:
    python evals/generate_eval_kb.py [target_dir]   # default: /tmp/nimbus-kb

The eval protocol and the question set live in docs/dev/rag-quality-eval.md.
"""

import sys
from pathlib import Path

import docx
import openpyxl


def _esc(s: str) -> str:
    return s.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")


def _make_pdf(pages_lines: list[list[str]]) -> bytes:
    """Minimal multi-page, multi-line, WinAnsi text-layer PDF (no deps)."""
    objects: list[bytes] = []
    n = len(pages_lines)
    page_objs = [4 + 2 * i for i in range(n)]
    kids = " ".join(f"{p} 0 R" for p in page_objs)
    objects.append(b"<< /Type /Catalog /Pages 2 0 R >>")
    objects.append(f"<< /Type /Pages /Kids [{kids}] /Count {n} >>".encode())
    objects.append(
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica"
        b" /Encoding /WinAnsiEncoding >>"
    )
    for i, lines in enumerate(pages_lines):
        body = " T* ".join(f"({_esc(line)}) Tj" for line in lines)
        content = f"BT /F1 11 Tf 14 TL 60 760 Td {body} ET".encode("latin-1")
        objects.append(
            (
                f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792]"
                f" /Resources << /Font << /F1 3 0 R >> >>"
                f" /Contents {page_objs[i] + 1} 0 R >>"
            ).encode()
        )
        objects.append(
            f"<< /Length {len(content)} >>\nstream\n".encode()
            + content
            + b"\nendstream"
        )
    out = bytearray(b"%PDF-1.4\n")
    offsets = []
    for num, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{num} 0 obj\n".encode() + body + b"\nendobj\n"
    xref = len(out)
    out += f"xref\n0 {len(objects) + 1}\n".encode() + b"0000000000 65535 f \n"
    for off in offsets:
        out += f"{off:010d} 00000 n \n".encode()
    out += (
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref}\n%%EOF\n"
    ).encode()
    return bytes(out)


def generate(base: Path) -> None:
    base.mkdir(parents=True, exist_ok=True)

    # ---- 1. Product guide (markdown, headed sections) --------------------
    (base / "guide-produit-nimbus.md").write_text(
        """# Guide produit Nimbus Analytics

## Présentation générale

Nimbus Analytics est une plateforme SaaS française d'analyse de données commerciales
destinée aux PME et ETI. Fondée en 2019 à Nantes, l'entreprise compte 47 salariés
et sert 312 clients actifs au 31 décembre 2025. La plateforme ingère les données de
vente (CRM, ERP, e-commerce) et produit des tableaux de bord prédictifs.

## Plans tarifaires

### Plan Starter

Le plan Starter coûte 89 euros HT par mois et par espace de travail. Il inclut
3 utilisateurs, 5 sources de données connectées, un historique de données de
13 mois et 50 000 lignes analysées par mois. Le support se fait par email
uniquement, avec un délai de réponse garanti de 48 heures ouvrées.

### Plan Business

Le plan Business coûte 290 euros HT par mois et par espace de travail. Il inclut
15 utilisateurs, 20 sources de données, un historique de 36 mois et 500 000 lignes
analysées par mois. Il ajoute les alertes automatiques, l'export programmé et
l'accès API en lecture (limité à 10 000 appels par jour). Support email et chat,
délai de réponse 8 heures ouvrées.

### Plan Enterprise

Le plan Enterprise est sur devis, à partir de 1 100 euros HT par mois. Utilisateurs
illimités, sources illimitées, historique de 60 mois, volume d'analyse négocié.
Il ajoute l'API complète en lecture-écriture (100 000 appels par jour), le SSO
(SAML 2.0 et OIDC), un environnement de préproduction dédié et un Customer Success
Manager nommé. Support prioritaire 24/7 avec délai de première réponse d'une heure.

## Limites techniques

La taille maximale d'un fichier importé manuellement est de 2 Go. Les connecteurs
natifs couvrent Salesforce, HubSpot, Shopify, PrestaShop, SAP Business One et
Odoo. Les rafraîchissements de données s'effectuent toutes les 4 heures sur le
plan Starter, toutes les heures sur le plan Business, et toutes les 15 minutes
sur le plan Enterprise.

## Module NimbusPredict

NimbusPredict est le module de prévision des ventes, disponible uniquement sur
les plans Business et Enterprise, en supplément de 120 euros HT par mois. Il
produit des prévisions à 3, 6 et 12 mois avec un intervalle de confiance affiché.
La précision moyenne constatée (MAPE) est de 11,4 % sur les prévisions à 3 mois.
""",
        encoding="utf-8",
    )

    # ---- 2. Client contract (docx: headings, clauses, pricing table) ------
    d = docx.Document()
    d.add_heading("Contrat cadre de services — Meridia Distribution SA", level=1)
    d.add_paragraph(
        "Contrat n° NB-2025-0147 conclu le 12 mai 2025 entre Nimbus Analytics SAS, "
        "immatriculée au RCS de Nantes sous le numéro 852 419 367, et Meridia "
        "Distribution SA, 14 avenue des Tilleuls, 69003 Lyon."
    )
    d.add_heading("Article 3 — Durée et renouvellement", level=2)
    d.add_paragraph(
        "Le présent contrat est conclu pour une durée initiale de trente-six (36) mois "
        "à compter du 1er juin 2025. Il se renouvelle ensuite tacitement par périodes "
        "de douze (12) mois, sauf dénonciation par l'une des parties avec un préavis "
        "de quatre-vingt-dix (90) jours avant l'échéance."
    )
    d.add_heading("Article 5 — Niveaux de service (SLA)", level=2)
    d.add_paragraph(
        "Le Prestataire garantit une disponibilité mensuelle de la Plateforme de "
        "99,7 %, mesurée hors fenêtres de maintenance programmées. Les fenêtres de "
        "maintenance sont limitées à six (6) heures par mois, notifiées au moins "
        "soixante-douze (72) heures à l'avance et planifiées entre 22h00 et 6h00."
    )
    d.add_heading("Article 6 — Pénalités", level=2)
    d.add_paragraph(
        "En cas de manquement au niveau de disponibilité garanti, le Client bénéficie "
        "d'un avoir égal à 5 % de la redevance mensuelle par tranche d'une heure "
        "d'indisponibilité au-delà du seuil, plafonné à 30 % de la redevance mensuelle. "
        "L'avoir doit être réclamé par écrit dans les trente (30) jours suivant l'incident."
    )
    d.add_heading("Article 9 — Responsabilité", level=2)
    d.add_paragraph(
        "La responsabilité totale cumulée du Prestataire, toutes causes confondues, "
        "est plafonnée au montant des redevances effectivement versées par le Client "
        "au cours des douze (12) derniers mois précédant le fait générateur."
    )
    d.add_heading("Conditions financières", level=2)
    t = d.add_table(rows=4, cols=3)
    for i, row in enumerate(
        [
            ("Poste", "Quantité", "Montant mensuel HT"),
            ("Plan Enterprise", "1 espace de travail", "1 450 €"),
            ("Module NimbusPredict", "1", "120 €"),
            ("Espaces additionnels", "2", "380 € (190 € l'unité)"),
        ]
    ):
        for j, val in enumerate(row):
            t.rows[i].cells[j].text = val
    d.add_paragraph(
        "Total mensuel HT : 1 950 €. Facturation trimestrielle à terme à échoir, "
        "paiement à 30 jours."
    )
    d.save(str(base / "contrat-cadre-meridia.docx"))

    # ---- 3. Financials (xlsx, 3 sheets) -----------------------------------
    wb = openpyxl.Workbook()
    ws1 = wb.active
    ws1.title = "CA trimestriel 2025"
    for row in [
        ["Trimestre", "CA total (k€)", "Dont abonnements (k€)", "Dont services (k€)", "Nouveaux clients"],
        ["T1 2025", 1240, 1085, 155, 18],
        ["T2 2025", 1378, 1196, 182, 24],
        ["T3 2025", 1456, 1301, 155, 21],
        ["T4 2025", 1689, 1478, 211, 33],
    ]:
        ws1.append(row)
    ws2 = wb.create_sheet("Charges 2025")
    for row in [
        ["Poste", "Montant annuel (k€)"],
        ["Masse salariale", 3120],
        ["Hébergement et infrastructure", 412],
        ["Marketing et acquisition", 587],
        ["Locaux et frais généraux", 296],
        ["R&D externalisée", 178],
    ]:
        ws2.append(row)
    ws3 = wb.create_sheet("Effectifs")
    for row in [
        ["Équipe", "Effectif au 31-12-2025"],
        ["Ingénierie", 21],
        ["Produit et design", 6],
        ["Ventes", 9],
        ["Customer Success", 7],
        ["Fonctions support", 4],
    ]:
        ws3.append(row)
    wb.save(str(base / "resultats-financiers-2025.xlsx"))

    # ---- 4. Security policy (3-page text-layer pdf) -----------------------
    page1 = [
        "Politique de securite des systemes d'information - Nimbus Analytics",
        "Version 4.2 - approuvee par le comite de direction le 3 fevrier 2025",
        "",
        "1. Hebergement et localisation des donnees",
        "L'ensemble des donnees clients est heberge chez OVHcloud, dans les",
        "datacenters de Gravelines (France) avec une replication synchrone vers",
        "le site de Roubaix. Aucune donnee n'est transferee hors de l'Union",
        "europeenne. Les sauvegardes completes sont realisees chaque nuit a 2h00",
        "et conservees quatre-vingt-dix (90) jours.",
        "",
        "2. Chiffrement",
        "Les donnees sont chiffrees au repos en AES-256 et en transit en TLS 1.3.",
        "Les cles de chiffrement sont gerees via un HSM dedie et font l'objet",
        "d'une rotation annuelle obligatoire.",
    ]
    page2 = [
        "3. Continuite d'activite",
        "L'objectif de point de reprise (RPO) est de quinze (15) minutes et",
        "l'objectif de delai de reprise (RTO) est de quatre (4) heures. Le plan",
        "de reprise d'activite est teste deux fois par an, en mars et en octobre.",
        "",
        "4. Gestion des acces",
        "L'authentification multifacteur est obligatoire pour tous les comptes",
        "internes. Les acces aux environnements de production sont revus chaque",
        "trimestre. Les journaux d'acces et d'audit sont conserves dix-huit (18)",
        "mois dans un stockage immuable.",
        "",
        "5. Certifications et conformite",
        "Nimbus Analytics est certifiee ISO 27001 depuis novembre 2023 et a",
        "obtenu son attestation SOC 2 Type II en juin 2025. La conformite RGPD",
        "est documentee dans le registre des traitements, revu semestriellement.",
    ]
    page3 = [
        "6. Gestion des incidents de securite",
        "Tout incident de securite est qualifie sous une (1) heure par l'equipe",
        "d'astreinte. En cas de violation de donnees personnelles, le delegue a",
        "la protection des donnees notifie la CNIL sous soixante-douze (72)",
        "heures et les clients concernes sous quarante-huit (48) heures.",
        "",
        "7. Tests d'intrusion",
        "Un test d'intrusion externe est realise chaque annee par un cabinet",
        "independant (derniere campagne : septembre 2025, cabinet Synacktiv).",
        "Les vulnerabilites critiques doivent etre corrigees sous sept (7) jours,",
        "les vulnerabilites majeures sous trente (30) jours.",
    ]
    (base / "politique-securite.pdf").write_bytes(_make_pdf([page1, page2, page3]))

    # ---- 5. Support FAQ (markdown) ----------------------------------------
    (base / "faq-support.md").write_text(
        """# FAQ Support Nimbus Analytics

## Quels sont les délais de réponse du support ?

Plan Starter : 48 heures ouvrées, par email uniquement. Plan Business : 8 heures
ouvrées, par email et chat. Plan Enterprise : première réponse sous 1 heure,
24/7, tous canaux y compris téléphone.

## Comment escalader un ticket bloquant ?

Pour les clients Enterprise, tout incident de priorité P1 (plateforme
inaccessible ou perte de données) peut être escaladé directement auprès du
Customer Success Manager nommé. En dehors des heures ouvrées, la ligne
d'astreinte est joignable au +33 2 85 52 41 90.

## Puis-je changer de plan en cours d'abonnement ?

L'upgrade est possible à tout moment et prend effet immédiatement, avec
facturation au prorata. Le downgrade prend effet à la prochaine échéance
annuelle et doit être demandé au moins 30 jours avant celle-ci.

## La plateforme est-elle accessible hors ligne ?

Non. Nimbus Analytics est une application 100 % cloud. Les exports PDF et
Excel programmés permettent toutefois de consulter les rapports hors
connexion.

## Comment sont gérées les données après résiliation ?

À la fin du contrat, les données du client restent exportables pendant
soixante (60) jours, puis sont supprimées définitivement des systèmes de
production et des sauvegardes dans un délai maximal de cent vingt (120) jours.
""",
        encoding="utf-8",
    )

    # ---- 6. Strategy committee notes (docx) -------------------------------
    d = docx.Document()
    d.add_heading("Comité stratégie — synthèse du 9 janvier 2026", level=1)
    d.add_paragraph(
        "Participants : direction générale, VP Produit, VP Ventes, DAF. "
        "Diffusion restreinte."
    )
    d.add_heading("Objectifs 2026", level=2)
    d.add_paragraph(
        "L'objectif d'ARR au 31 décembre 2026 est fixé à 8,5 millions d'euros, contre "
        "6,2 millions d'euros constatés fin 2025, soit une croissance cible de 37 %. "
        "Le taux d'attrition (churn brut) devra être maintenu sous 9 % annuel."
    )
    d.add_heading("Lancements produit", level=2)
    d.add_paragraph(
        "Le module NimbusPredict version 2, intégrant la prise en compte de la "
        "saisonnalité multi-magasins, est planifié pour le deuxième trimestre 2026. "
        "Le connecteur Microsoft Dynamics 365 sortira au premier trimestre 2026. "
        "Un assistant conversationnel d'exploration des données (projet interne "
        "« Cumulus ») entrera en bêta privée au quatrième trimestre 2026."
    )
    d.add_heading("Expansion internationale", level=2)
    d.add_paragraph(
        "Ouverture d'une filiale en Allemagne (Munich) au second semestre 2026, avec "
        "un objectif de 3 recrutements commerciaux et une localisation complète de la "
        "plateforme en allemand. Le marché espagnol est reporté à 2027."
    )
    d.add_heading("Risques identifiés", level=2)
    d.add_paragraph(
        "Dépendance au connecteur Salesforce (38 % des clients) ; hausse des coûts "
        "d'hébergement estimée à 12 % ; tension sur le recrutement d'ingénieurs data. "
        "Budget de couverture des risques provisionné à 250 000 euros."
    )
    d.save(str(base / "notes-comite-strategie.docx"))

    print(f"Corpus Nimbus generated in {base} ({len(list(base.iterdir()))} files)")


if __name__ == "__main__":
    generate(Path(sys.argv[1]) if len(sys.argv) > 1 else Path("/tmp/nimbus-kb"))
